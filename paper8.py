import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import win32print
import win32ui
import os
import threading

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.recording = False  # Flag to indicate recording status
        self.create_menu()
        self.create_toolbar()
        self.create_textbox()
        self.preview_window = None
        self.recordings = []

    def create_menu(self):
        menubar = tk.Menu(self.root)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_command(label="Exit", command=self.quit_app)
        menubar.add_cascade(label="File", menu=file_menu)
        
        predefined_menu = tk.Menu(menubar, tearoff=0)
        predefined_menu.add_command(label="Text 1", command=lambda: self.insert_predefined_text("text1.txt"))
        predefined_menu.add_command(label="Text 2", command=lambda: self.insert_predefined_text("text2.txt"))
        predefined_menu.add_command(label="Text 3", command=lambda: self.insert_predefined_text("text3.txt"))
        # Add more predefined text options as needed
        menubar.add_cascade(label="Text", menu=predefined_menu)
        
        self.root.config(menu=menubar)

    def create_toolbar(self):
        toolbar = tk.Frame(self.root, bg="lightgray")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_button = tk.Button(toolbar, text="Start Recording", command=self.start_recording)
        self.recording_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.LEFT, padx=5, pady=5)

    def create_textbox(self):
        self.textbox = tk.Text(self.root, height=20, width=60)
        self.textbox.pack(pady=10)

    def save_to_docx(self):
        if self.recordings:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            file_path = os.path.join(current_dir, "speech_to_text.docx")
            try:
                doc = Document(file_path) if os.path.exists(file_path) else Document()
                for text in self.recordings:
                    doc.add_paragraph(text)
                doc.save(file_path)
                self.recordings.clear()
            except Exception as e:
                print("Error saving file:", e)

    def start_recording(self):
        if not self.recording:
            self.recording = True
            self.recording_button.config(text="Stop Recording")
            self.status_label.config(text="Recording...", fg="red")
            threading.Thread(target=self._recording_thread).start()
        else:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread(self):
        options = []  # Initialize a list to store the options

        with self.microphone as source:
            print("Speak your multiple-choice question in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append the question text to the options list
            options.append(text.strip())

            # Prompt the user to speak each option
            for i in range(4):
                print(f"Speak option {chr(65+i)} in Hindi:")
                with self.microphone as source:
                    self.recognizer.adjust_for_ambient_noise(source, duration=1)
                    audio = self.recognizer.listen(source)
                option_text = self.recognizer.recognize_google(audio, language="hi-IN")
                print(f"You said option {chr(65+i)}:", option_text)
                options.append(option_text.strip())

            # Format the recognized text as a multiple-choice question with options
            mcq_text = f"{self.counter}. {options[0]} ?\nA. {options[1]}\nB. {options[2]}\nC. {options[3]}\nD. {options[4]}\n"
            self.counter += 1  # Increment counter

            self.textbox.insert(tk.END, mcq_text + "\n")  # Display in text box
            self.recordings.append(mcq_text)  # Save for batch saving

        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def insert_predefined_text(self, filename=None):
        try:
            if filename:
                file_path = os.path.join(os.path.dirname(__file__), filename)
                print("File path:", file_path)  # Print the absolute file path for debugging
                with open(file_path, "r", encoding="utf-8") as file:
                    predefined_text = file.read()
                    self.textbox.insert(tk.END, predefined_text + "\n\n")  # Insert predefined text
                    
                    self.recordings.append(predefined_text)  # Save for batch saving
        except FileNotFoundError:
            print("Predefined text file not found.")

    def save_text(self):
        threading.Thread(target=self.save_to_docx).start()

    def preview_text(self):
        if self.preview_window is None or not self.preview_window.winfo_exists():
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title("Preview")
            preview_textbox = tk.Text(self.preview_window, height=20, width=60)
            preview_textbox.pack(pady=10)
            preview_text = self.textbox.get("1.0", tk.END)
            preview_textbox.insert(tk.END, preview_text)
        else:
            self.preview_window.deiconify()
            self.preview_window.lift()

    def print_text(self):
        text = self.textbox.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)

    def quit_app(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()

if __name__ == "__main__":
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
