import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import RGBColor
import win32print
import win32ui
import os
import threading
import time
import subprocess

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Exam Paper Creater")
        self.root.geometry("800x630")
        self.root.resizable(False, False)
        #self.root.state("zoomed")  # Maximize the window    
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.counter4opt = 0  # Initialize counter for numbering        
        self.hindi_numbers = ['(क)', '(ख)', '(ग)', '(घ)', '(ङ)', '(च)', '(छ)', '(ज)', '(झ)', '(ञ)', '(ट)', '(ठ)', '(ड)', '(ढ)', '(ण)', '(त)', '(थ)', '(द)', '(ध)', '(न)', '(प)', '(फ)', '(ब)', '(भ)', '(म)', '(य)', '(र)', '(ल)', '(व)', '(श)', '(ष)', '(स)', '(ह)', '(अ)', '(आ)', '(इ)', '(ई)', '(उ)', '(ऊ)', '(ऋ)', '(ए)', '(ऐ)', '(ओ)', '(औ)', '(अं)', '(अः)', '(क्ष)', '(त्र)', '(ज्ञ)']  # Hindi numbers
        self.counter2 = 1  # Initialize counter for numbering
        self.roman_numbers = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI', 'XVII', 'XVIII', 'XIX','XX']  # Hindi numbers
        self.selected_punctuation = "?"  # Default punctuation
        self.recording = False  # Flag to indicate recording status
        self.MCQ_numbers = ['  (i)','(ii)','(iii)','(iv)']
        self.roman2_numbers = ['  (i)','(ii)','(iii)','(iv)', '(v)', '(vi)']

        # Initialize counter for numbering
        self.counter4opt = 0
        
        self.another_predefined_texts = {
            "Short Type": "This is the text for Example first.",
            "Long Type": "This is the text for Example second.",
            "Multiple Choice": "This is the text for Example third."
        }        
        
        self.predefined_texts = {
            "Yearly exam": "समय - 						        वार्षिक परीक्षा",
            "Half Yearly exam": "समय - 						        अर्ध-वार्षिक परीक्षा",
            "The note": "                                                                                   नोट – सभी प्रश्न अनिवार्य है –"
        }
        
        self.my_predefined_texts = {
            "Class-1": {
                "Hindi": "                                                                                     कक्षा – 1 विषय – हिन्दी",
                "Social": "                                                                                     कक्षा – 1 विषय – सामाजिक",
                "Science": "                                                                                     कक्षा – 1 विषय – विज्ञान"
            },
            "Class-2": {
                "Hindi": "                                                                                     कक्षा – 2 विषय – हिन्दी",
                "Social": "                                                                                     कक्षा – 2 विषय – सामाजिक",
                "Science": "                                                                                     कक्षा – 2 विषय – विज्ञान"
            },
            "Class-3": {
                "Hindi": "                                                                                     कक्षा – 3 विषय – हिन्दी",
                "Social": "                                                                                     कक्षा – 3 विषय – सामाजिक",
                "Science": "                                                                                     कक्षा – 3 विषय – विज्ञान"
            }
        } 

 
        self.create_widgets()
        self.preview_window = None
        self.recordings = []

    def create_widgets(self):
        # Menu
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)
        
        # file menu
        file_menu = tk.Menu(menu, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Open Saved File", command=self.open_saved_file)  # Add command to open saved file
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_separator()                
        file_menu.add_command(label="Exit", command=self.quit_app)
        menu.add_cascade(label="File", menu=file_menu)
        
        
        #Exam title
        text_menu2 = tk.Menu(menu, tearoff=0)
        for label in self.predefined_texts:
            text_menu2.add_command(label=label, command=lambda txt=self.predefined_texts[label]: self.insert_predefined_text2(txt))
        menu.add_cascade(label="ExamTitle", menu=text_menu2)   

        # "Class" submenu
        many_text_submenu = tk.Menu(menu, tearoff=0)
        for name, texts in self.my_predefined_texts.items():
            submenu = tk.Menu(many_text_submenu, tearoff=0)
            for label, text in texts.items():
                submenu.add_command(label=label, command=lambda txt=text: self.insert_predefined_text2(txt))
            many_text_submenu.add_cascade(label=name, menu=submenu)
        menu.add_cascade(label="Class", menu=many_text_submenu)

        
        #Questions type
        another_predefined_submenu = tk.Menu(menu, tearoff=0)
        for label, text in self.another_predefined_texts.items():
            another_predefined_submenu.add_command(label=label, command=lambda txt=text: self.insert_predefined_text3(txt))
        menu.add_cascade(label="Questions", menu=another_predefined_submenu)   


        #My text
        text_menu = tk.Menu(menu, tearoff=0)
        text_menu.add_command(label="Text 1", command=lambda: self.insert_predefined_text("text1.txt"))
        text_menu.add_command(label="Text 2", command=lambda: self.insert_predefined_text("text2.txt"))
        text_menu.add_command(label="Text 3", command=lambda: self.insert_predefined_text("text3.txt"))
        menu.add_cascade(label="MyText", menu=text_menu)
        
    
        

     





        # Toolbar
        toolbar = tk.Frame(self.root, bg="#f0f0f0")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_button = tk.Button(toolbar, text="Insert Questions", command=self.start_recording, state="normal")
        self.recording_button.pack(side=tk.LEFT, padx=5, pady=5)
        self.recording_button.configure(background='#f5e1bf')          
        remove_last_button = tk.Button(toolbar, text=" X ", command=self.remove_ques) #for question
        remove_last_button.pack(side=tk.LEFT, padx=5, pady=5)
        remove_last_button.configure(background='#f5e1bf')  
        
        
        self.recording_button2 = tk.Button(toolbar, text="Insert Options", command=self.start_recording2, state="normal")
        self.recording_button2.pack(side=tk.LEFT, padx=5, pady=5)       
        self.recording_button2.configure(background='#c4d2ff')         
        self.decrease_button = tk.Button(toolbar, text=" X ", command=self.remove_mcq) #for mcqs
        self.decrease_button.pack(side=tk.LEFT, padx=5, pady=5)
        self.decrease_button.configure(background='#c4d2ff')  

        self.recording_button3 = tk.Button(toolbar, text="Insert Sentence", command=self.start_recording3, state="normal")
        self.recording_button3.pack(side=tk.LEFT, padx=5, pady=5)  
        self.recording_button3.configure(background='#f5bcde')          
        self.rem_sen_button = tk.Button(toolbar, text=" X ", command=self.remove_sentence) #for mcqs
        self.rem_sen_button.pack(side=tk.LEFT, padx=5, pady=5)
        self.rem_sen_button.configure(background='#f5bcde')        
        
        self.recording_button4 = tk.Button(toolbar, text="Insert List", command=self.start_recording4, state="normal")
        self.recording_button4.pack(side=tk.LEFT, padx=5, pady=5)  
        self.recording_button4.configure(background='#bbdef2')          
        self.rem_list_button = tk.Button(toolbar, text=" X ", command=self.remove_list) #for mcqs
        self.rem_list_button.pack(side=tk.LEFT, padx=5, pady=5)
        self.rem_list_button.configure(background='#bbdef2')             
        
        # -----------RIGHT-------------     
        
        # Button to insert blank line
        blank_line_button = tk.Button(toolbar, text="New-Line", command=self.insert_blank_line)
        blank_line_button.pack(side=tk.RIGHT, padx=5, pady=5)
        blank_line_button.configure(background='#c9c7c9')                 


        blank_line_button = tk.Button(toolbar, text="Box", command=self.insert_box)
        blank_line_button.pack(side=tk.RIGHT, padx=5, pady=5)
        blank_line_button.configure(background='#c9c7c9')          
        

        blank_line_button = tk.Button(toolbar, text="Blank", command=self.insert_blank)
        blank_line_button.pack(side=tk.RIGHT, padx=5, pady=5)
        blank_line_button.configure(background='#c9c7c9')   
        
 
        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.RIGHT, padx=5, pady=5)

 



        # Text box
        self.text_box = tk.Text(self.root, font=("Arial", 14), wrap="word")
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Bind right-click menu
        self.text_box.bind("<Button-3>", self.show_context_menu)


    def maximize_window(self):
        self.root.state("zoomed")  # Maximize the window    

    def restore_window(self):
        self.root.state("normal")  # Restore the window to its normal state    

    def resize_window(self):
        self.root.resizable(True, True)

    def insert_blank_line(self):
        """
        Inserts a blank line after the current text position in the text box.
        """
        self.text_box.insert(tk.END, "\n", "silver_line")
        self.text_box.tag_configure("silver_line", background="#C0C0C0")
        self.text_box.see(tk.END)  # Scrolls the text box to the end after insertion


    def insert_blank(self):
        self.text_box.insert(tk.INSERT, "______")

    def insert_box(self):
        self.text_box.insert(tk.INSERT, " [   ]")


    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Select All", command=self.select_all)
        context_menu.add_command(label="Copy All", command=self.copy_text)
        context_menu.add_separator()
        context_menu.add_command(label="Copy", command=self.copy_selected_text)
        context_menu.add_command(label="Cut", command=self.cut_selected_text)
        context_menu.add_command(label="Paste", command=self.paste_text)
        context_menu.add_separator()        
        context_menu.add_command(label="Punctuation: Period", command=lambda: self.set_punctuation("|"))
        context_menu.add_command(label="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))  
        context_menu.add_separator()
        context_menu.add_command(label="Maximize", command=self.maximize_window)
        context_menu.add_command(label="Restore", command=self.restore_window)
        context_menu.add_command(label="I want to resize the window", command=self.resize_window)        
        context_menu.tk_popup(event.x_root, event.y_root)

    def copy_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(text_content)
        messagebox.showinfo("Success", "Text copied successfully!")

    def save_to_docx(self):
        if self.recordings:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            file_path = os.path.join(current_dir, "speech_to_text.docx")
            try:
                doc = Document(file_path) if os.path.exists(file_path) else Document()
                for text in self.recordings:
                    doc.add_paragraph(text)
                doc.save(file_path)
                self.recordings.clear()
            except Exception as e:
                print("Error saving file:", e)

    def start_recording(self):
        if not self.recording:
            self.recording = True
            self.recording_button.config(text="Stop Inserting")
            self.status_label.config(text="Listening...", fg="red")
            threading.Thread(target=self._recording_thread).start()
        else:
            self.recording = False
            self.recording_button.config(text="Insert Questions")
            self.status_label.config(text="Ready", fg="green")
            
    def start_recording2(self):
        if not self.recording:
            self.recording = True
            self.recording_button2.config(text="Stop Inserting")
            self.status_label.config(text="Listening...", fg="red")
            threading.Thread(target=self._recording_thread2).start()
        else:
            self.recording = False
            self.recording_button2.config(text="Insert Options")
            self.status_label.config(text="Ready", fg="green")            
 
    def start_recording3(self):
        if not self.recording:
            self.recording = True
            self.recording_button3.config(text="Stop Inserting")
            self.status_label.config(text="Listening...", fg="red")
            threading.Thread(target=self._recording_thread3).start()
        else:
            self.recording = False
            self.recording_button3.config(text="Insert Sentence")
            self.status_label.config(text="Ready", fg="green")   

    def start_recording4(self):
        if not self.recording:
            self.recording = True
            self.recording_button4.config(text="Stop Inserting")
            self.status_label.config(text="Listening...", fg="red")
            threading.Thread(target=self._recording_thread4).start()
        else:
            self.recording = False
            self.recording_button4.config(text="Insert List")
            self.status_label.config(text="Ready", fg="green")    

    def _recording_thread(self):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.hindi_numbers[self.counter - 1]}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            # Insert voice-recognized text without bold formatting
            self.text_box.insert(tk.END, numbered_text + "\n")
            self.recordings.append(numbered_text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
            messagebox.showinfo("Error", "Please connect your device to Internet!")
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button.config(text="Insert Questions")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread2(self):
        print("Counter value at the start:", self.counter4opt)  # Display initial counter value
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text and increment the counter
            numbered_text = f"{self.MCQ_numbers[self.counter4opt]} {text}"
            # Insert voice-recognized text without bold formatting
            self.text_box.insert(tk.END, numbered_text + "    ")
            self.recordings.append(numbered_text)  # Save for batch saving

            # Increment counter and check if it reaches the end
            self.counter4opt += 1
            if self.counter4opt == len(self.MCQ_numbers):
                # Start a new line and reset the counter
                self.text_box.insert(tk.END, "\n\n")
                self.counter4opt = 0  # Reset the counter to 0
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
            messagebox.showerror("error", "Could not request results from Google Speech Recognition service; {0}".format(e))
            messagebox.showinfo("Error", "Please connect your device to Internet!")
        except Exception as e:
            print("Error:", e)
            messagebox.showerror("Error", e)
        finally:
            self.recording = False
            self.recording_button2.config(text="Insert Options")
            self.status_label.config(text="Ready", fg="green")
            

    def _recording_thread3(self):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Insert voice-recognized text into the text box
            self.text_box.insert(tk.END, text + "\n")
            self.recordings.append(text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
            messagebox.showinfo("Error", "Please connect your device to Internet!")
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button3.config(text="Insert Sentence")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread4(self):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.roman2_numbers[self.counter - 1]}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            # Insert voice-recognized text without bold formatting
            self.text_box.insert(tk.END, numbered_text + "\n")
            self.recordings.append(numbered_text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
            messagebox.showinfo("Error", "Please connect your device to Internet!")
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button4.config(text="Insert List")
            self.status_label.config(text="Ready", fg="green")

    
    # for mcqs
    def remove_mcq(self):
        # Find the start position of the last MCQ numbering
        start_pos = self.text_box.search(f"\\({self.MCQ_numbers[self.counter4opt - 1]}\\)", "1.0", stopindex=tk.END, regexp=True)
        if start_pos:
            # Find the end position of the last MCQ numbering
            end_pos = self.text_box.search(f"\\({self.MCQ_numbers[self.counter4opt]}\\)", start_pos, stopindex=tk.END, regexp=True)
            if not end_pos:
                end_pos = self.text_box.search("\n", start_pos, stopindex=tk.END)
            self.counter4opt -= 1
            # Check if the start position is at the beginning of the text
            if start_pos != "1.0":
                start_pos += "-1c"  # Adjust the start position to include the previous character (newline)
            self.text_box.delete(start_pos, end_pos + " lineend")








    
    
    #for questions
    def remove_ques(self):
        if self.counter > 1:
            # Find the start position of the last Hindi numbering
            start_pos = self.text_box.search(f"\\({self.hindi_numbers[self.counter - 2]}\\)", "1.0", stopindex=tk.END, regexp=True)
            if start_pos:
                # Find the end position of the last Hindi numbering
                end_line = int(start_pos.split('.')[0]) + 1  # Get the line number of the next Hindi numbering
                end_pos = f"{end_line}.0"  # Start of the next Hindi numbering
                self.counter -= 1
                self.text_box.delete(start_pos, end_pos)


    def remove_sentence(self):
        start_pos = self.text_box.search("You said:", "1.0", stopindex=tk.END, regexp=True)
        if start_pos:
            end_pos = self.text_box.search("\n", start_pos, tk.END, regexp=False)
            if end_pos:
                self.text_box.delete(start_pos, end_pos + "+1c")


    #for list
    def remove_list(self):
        if self.counter > 1:
            # Find the start position of the last Hindi numbering
            start_pos = self.text_box.search(f"\\({self.roman2_numbers[self.counter - 2]}\\)", "1.0", stopindex=tk.END, regexp=True)
            if start_pos:
                # Find the end position of the last Hindi numbering
                end_line = int(start_pos.split('.')[0]) + 1  # Get the line number of the next Hindi numbering
                end_pos = f"{end_line}.0"  # Start of the next Hindi numbering
                self.counter -= 1
                self.text_box.delete(start_pos, end_pos)

    def set_punctuation(self, punctuation):
        self.selected_punctuation = punctuation

    def save_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        default_file_name = "My Exam Paper"  # Default file name without extension
        file_name = simpledialog.askstring("Save As", "Enter file name (without extension):", initialvalue=default_file_name)
        if file_name:
            file_name += ".docx"  # Append .docx extension
            file_path = os.path.join(os.getcwd(), file_name)
            try:
                doc = Document()  # Create a new Document object

                # Set page size to A4
                section = doc.sections[0]
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)

                # Set margins to narrow
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Define paragraph style with minimum spacing
                style = doc.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                # Iterate over each line in the text content
                for line in text_content.split("\n"):
                    # Check if the line contains a predefined text
                    predefined_text_found = False
                    
                    # Check in the first dictionary
                    for predefined_text in self.predefined_texts.values():
                        if predefined_text in line:
                            # Apply bold formatting to the entire paragraph
                            paragraph = doc.add_paragraph(line)
                            paragraph.style = 'Strong'
                            predefined_text_found = True
                            break
                    
                    # If not found in the first dictionary, check in the second dictionary
                    if not predefined_text_found:
                        for predefined_text in self.another_predefined_texts.values():
                            if predefined_text in line:
                                # Apply bold formatting to the entire paragraph
                                paragraph = doc.add_paragraph(line)
                                paragraph.style = 'Strong'
                                predefined_text_found = True
                                break
                    
                    # If not found in the second dictionary, check in the third dictionary
                    if not predefined_text_found:
                        for person, person_texts in self.my_predefined_texts.items():
                            for text_name, text_content in person_texts.items():
                                if text_content in line:
                                    # Apply bold formatting to the entire paragraph
                                    paragraph = doc.add_paragraph(line)
                                    paragraph.style = 'Strong'
                                    predefined_text_found = True
                                    break
                            if predefined_text_found:
                                break
                    
                    if not predefined_text_found:
                        # If the line does not contain a predefined text, add it normally
                        paragraph = doc.add_paragraph(line, style)
                        # Apply font type and size to the entire paragraph
                        for run in paragraph.runs:
                            run.font.name = "Arial"  # Set font type to Arial
                            run.font.size = Pt(11)  # Set font size to 14 points

                doc.save(file_path)  # Save the document
                messagebox.showinfo("Success", f"Text saved successfully as '{file_name}'!")
            except PermissionError as e:
                messagebox.showerror("Error", f"Please close the {file_name} file before saving.")
            except Exception as e:
                print("Error saving file:", e)


    def insert_predefined_text(self, filename=None):
        self.counter = 1  # Reset the counter to 1
        try:
            if filename:
                file_path = os.path.join(os.path.dirname(__file__) + "\\myText", filename)
                print("File path:", file_path)  # Print the absolute file path for debugging
                with open(file_path, "r", encoding="utf-8") as file:
                    predefined_text = file.readlines()
                    numbered_text = ""
                    for line in predefined_text:
                        # Apply bold formatting to each line
                        numbered_text += f"\n{self.roman_numbers[self.counter2 - 1]}. {line.strip()}\n"
                        self.counter2 += 1
                    self.text_box.insert(tk.END, numbered_text)  # Insert predefined text
                    
                    # Apply bold formatting to the entire inserted text
                    self.text_box.tag_configure("bold", font=("Arial", 14, "bold"))
                    self.text_box.tag_add("bold", "end - 2 lines", "end - 1 lines")
                    
                    self.recordings = [numbered_text]  # Reset the recordings list
        except FileNotFoundError:
            print("Predefined text file not found.")


    def insert_predefined_text2(self, text=None):
        if text is None:
            text = self.predefined_texts["Yearly exam"]  # Default to the first predefined text
        
        # Apply bold formatting to the predefined text
        self.text_box.insert(tk.END, text + "\n", "predefined")  
        self.text_box.tag_configure("predefined", font=("Arial", 14, "bold"))


    def insert_predefined_text3(self, text=None):
        self.counter = 1  # Reset the counter to 1
        if text is None:
            text = self.predefined_texts["Yearly exam"]  # Default to the first predefined text

        # Split the text into sentences
        sentences = text.split(". ")

        # Add Roman numbering to each sentence
        formatted_text = ""
        for index, sentence in enumerate(sentences, start=self.counter2):
            formatted_text += f"{self.roman_numbers[index - 1]}. {sentence} "

        # Increment the counter for numbering
        self.counter2 += len(sentences)

        # Apply bold formatting to the predefined text
        self.text_box.insert(tk.END, "\n" + formatted_text + "\n", "predefined")
        self.text_box.tag_configure("predefined", font=("Arial", 14, "bold"))

 
    def save_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        default_file_name = "My Exam Paper"  # Default file name without extension
        file_name = simpledialog.askstring("Save As", "Enter file name (without extension):", initialvalue=default_file_name)
        if file_name:
            file_name += ".docx"  # Append .docx extension
            file_path = os.path.join(os.getcwd(), file_name)
            try:
                doc = Document()  # Create a new Document object

                # Set page size to A4
                section = doc.sections[0]
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)

                # Set margins to narrow
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Define paragraph style with minimum spacing
                style = doc.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                # Iterate over each line in the text content
                for line in text_content.split("\n"):
                    # Check if the line contains a predefined text
                    predefined_text_found = False
                    
                    # Check in the first dictionary
                    for predefined_text in self.predefined_texts.values():
                        if predefined_text in line:
                            # Apply bold formatting to the line
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run(line)
                            run.font.bold = True
                            run.font.size = Pt(12)  # Set font size to 14 points
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                            predefined_text_found = True
                            break
                    
                    # If not found in the first dictionary, check in the second dictionary
                    if not predefined_text_found:
                        for predefined_text in self.another_predefined_texts.values():
                            if predefined_text in line:
                                # Apply bold formatting to the line
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(line)
                                run.font.bold = True
                                run.font.size = Pt(12)  # Set font size to 14 points
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                                predefined_text_found = True
                                break
                    
                    # If not found in the second dictionary, check in the third dictionary
                    if not predefined_text_found:
                        for person, person_texts in self.my_predefined_texts.items():
                            for text_name, text_content in person_texts.items():
                                if text_content in line:
                                    # Apply bold formatting to the line
                                    paragraph = doc.add_paragraph()
                                    run = paragraph.add_run(line)
                                    run.font.bold = True
                                    run.font.size = Pt(12)  # Set font size to 14 points
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                                    predefined_text_found = True
                                    break
                            if predefined_text_found:
                                break
                    
                    if not predefined_text_found:
                        # If the line does not contain a predefined text, add it normally
                        paragraph = doc.add_paragraph(line, style)
                        # Apply font type and size to the entire paragraph
                        for run in paragraph.runs:
                            run.font.name = "Arial"  # Set font type to Arial
                            run.font.size = Pt(11)  # Set font size to 14 points

                doc.save(file_path)  # Save the document
                messagebox.showinfo("Success", f"Text saved successfully as '{file_name}'!")
            except PermissionError as e:
                messagebox.showerror("Error", f"Please close the {file_name} file before saving.")
            except Exception as e:
                print("Error saving file:", e)




    def open_saved_file(self):
        current_dir = os.getcwd()
        file_path = filedialog.askopenfilename(initialdir=current_dir, title="Select file", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
        if file_path:
            try:
                os.startfile(file_path)  # Open the file using default application
            except Exception as e:
                print("Error opening file:", e)



    def preview_text(self):
        if self.preview_window is None or not self.preview_window.winfo_exists():
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title("Preview")
            preview_textbox = tk.Text(self.preview_window, font=("Arial", 14), wrap="word")
            preview_textbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            preview_text = self.text_box.get("1.0", tk.END)
            preview_textbox.insert(tk.END, preview_text)
        else:
            self.preview_window.deiconify()
            self.preview_window.lift()

    def print_text(self):
        text = self.text_box.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)


                

    def select_all(self):
        self.text_box.tag_add(tk.SEL, "1.0", tk.END)
        self.text_box.mark_set(tk.INSERT, "1.0")
        self.text_box.see(tk.INSERT)
        return 'break'

    def copy_selected_text(self):
        self.text_box.clipboard_clear()
        text = self.text_box.get(tk.SEL_FIRST, tk.SEL_LAST)
        self.text_box.clipboard_append(text)

    def cut_selected_text(self):
        self.copy_selected_text()
        self.text_box.delete(tk.SEL_FIRST, tk.SEL_LAST)

    def paste_text(self):
        text = self.text_box.clipboard_get()
        self.text_box.insert(tk.INSERT, text)

    def quit_app(self):
        self.save_to_docx()
        self.root.quit()                

    def quit_app(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()

if __name__ == "__main__":

    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
