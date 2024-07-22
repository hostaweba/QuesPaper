import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import win32print
import win32ui
import os
import threading

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.root.geometry("800x600")
        self.root.resizable(False, False)
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.hindi_numbers = ['(i)', '(ii)', '(iii)', '(iv)']  # Hindi numbers
        self.roman_numbers = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']  # Roman numerals
        self.selected_punctuation = "?"  # Default punctuation
        self.recording = False  # Flag to indicate recording status
        self.predefined_text_dict = {  # Predefined text dictionary
            'I': 'First predefined text',
            'II': 'Second predefined text',
            'III': 'Third predefined text'
        }
        self.create_widgets()
        self.preview_window = None
        self.recordings = []

    def create_widgets(self):
        # Menu
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.quit_app)
        menu.add_cascade(label="File", menu=file_menu)

        text_menu = tk.Menu(menu, tearoff=0)
        for i in range(len(self.roman_numbers)):
            text_menu.add_command(label=f"Text {self.roman_numbers[i]}", command=lambda index=i: self.insert_predefined_text(index))
        menu.add_cascade(label="Text", menu=text_menu)

        # Toolbar
        toolbar = tk.Frame(self.root, bg="#f0f0f0")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_buttons = []
        for i in range(4):
            button = tk.Button(toolbar, text=f"({chr(0x0968+i)}) Start", command=lambda index=i: self.start_recording(index))
            button.pack(side=tk.LEFT, padx=5, pady=5)
            self.recording_buttons.append(button)

        punctuation_period_button = tk.Button(toolbar, text="Punctuation: Period", command=lambda: self.set_punctuation("."))
        punctuation_period_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_question_button = tk.Button(toolbar, text="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))
        punctuation_question_button.pack(side=tk.LEFT, padx=5, pady=5)

        remove_last_button = tk.Button(toolbar, text="Remove number", command=self.remove_number)
        remove_last_button.pack(side=tk.LEFT, padx=5, pady=5)

        copy_button = tk.Button(toolbar, text="Copy All", command=self.copy_text)
        copy_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.LEFT, padx=5, pady=5)

        # Text box
        self.text_box = tk.Text(self.root, font=("Arial", 14), wrap="word")
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Bind right-click menu
        self.text_box.bind("<Button-3>", self.show_context_menu)

    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Select All", command=self.select_all)
        context_menu.add_separator()
        context_menu.add_command(label="Copy", command=self.copy_selected_text)
        context_menu.add_command(label="Cut", command=self.cut_selected_text)
        context_menu.add_command(label="Paste", command=self.paste_text)
        context_menu.tk_popup(event.x_root, event.y_root)

    def copy_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(text_content)
        messagebox.showinfo("Success", "Text copied successfully!")

    def save_to_docx(self):
        if self.recordings:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            file_path = os.path.join(current_dir, "speech_to_text.docx")
            try:
                doc = Document(file_path) if os.path.exists(file_path) else Document()
                for text in self.recordings:
                    doc.add_paragraph(text)
                doc.save(file_path)
                self.recordings.clear()
            except Exception as e:
                print("Error saving file:", e)

    def start_recording(self, index):
        if not self.recording:
            self.recording = True
            self.recording_buttons[index].config(text=f"({chr(0x0968+index)}) Stop")
            self.status_label.config(text="Recording...", fg="red")
            threading.Thread(target=self._recording_thread, args=(index,)).start()
        else:
            self.recording = False
            self.recording_buttons[index].config(text=f"({chr(0x0968+index)}) Start")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread(self, index):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.hindi_numbers[self.counter - 1]}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            self.text_box.insert(tk.END, numbered_text + "    ")  # Display in text box
            self.recordings.append(numbered_text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_buttons[index].config(text=f"({chr(0x0968+index)}) Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def set_punctuation(self, punctuation):
        self.selected_punctuation = punctuation

    def insert_predefined_text(self, index):
        try:
            key = self.roman_numbers[index]
            text = self.predefined_text_dict[key]
            numbered_text = f"{self.roman_numbers[index]}. {text}\n\n"
            self.text_box.insert(tk.END, numbered_text)  # Insert predefined text
            self.recordings = [numbered_text]  # Reset the recordings list
        except Exception as e:
            print("Error inserting predefined text:", e)


    def save_text(self):
        threading.Thread(target=self.save_to_docx).start()

    def preview_text(self):
        if self.preview_window is None or not self.preview_window.winfo_exists():
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title("Preview")
            preview_textbox = tk.Text(self.preview_window, font=("Arial", 14), wrap="word")
            preview_textbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            preview_text = self.text_box.get("1.0", tk.END)
            preview_textbox.insert(tk.END, preview_text)
        else:
            self.preview_window.deiconify()
            self.preview_window.lift()

    def print_text(self):
        text = self.text_box.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)

    def remove_number(self):
        if self.counter > 1:
            # Find the start position of the last line number
            start_pos = self.text_box.search(f"\\(.*?\\)", "1.0", stopindex=tk.END, regexp=True)
            print("Start position:", start_pos)  # Debugging print statement
            if start_pos:
                # Find the end position of the last line
                end_pos = f"{self.hindi_numbers[self.counter - 2]}.999"  # Use 999 as the column number to get the end of the line
                print("End position:", end_pos)  # Debugging print statement

                self.counter -= 1
                self.text_box.delete(start_pos, end_pos)

    def quit_app(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()

if __name__ == "__main__":
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
