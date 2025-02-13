import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
import win32print
import win32ui
import os
import threading
import time

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.root.geometry("800x600")
        self.root.resizable(False, False)
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.hindi_numbers = ['(क)', '(ख)', '(ग)', '(घ)', '(ङ)', '(च)', '(छ)', '(ज)', '(झ)', '(ञ)']  # Hindi numbers
        self.counter2 = 1  # Initialize counter for numbering
        self.roman_numbers = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI', 'XVII', 'XVIII', 'XIX','XX']  # Hindi numbers
        self.selected_punctuation = "?"  # Default punctuation
        self.recording = False  # Flag to indicate recording status
        self.predefined_texts = {
            "Yearly exam": "समय - 							वार्षिक परीक्षा",
            "Class": {
                "Class-1": {
                    "Eng": "Yearly exam: समय - 							वार्षिक परीक्षा\nEnglish text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "Hindi": "Yearly exam: समय - 							वार्षिक परीक्षा\nHindi text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "Science": "Yearly exam: समय - 							वार्षिक परीक्षा\nScience text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "Social": "Yearly exam: समय - 							वार्षिक परीक्षा\nSocial text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "Math": "Yearly exam: समय - 							वार्षिक परीक्षा\nMath text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "Drawing": "Yearly exam: समय - 							वार्षिक परीक्षा\nDrawing text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "History": "Yearly exam: समय - 							वार्षिक परीक्षा\nHistory text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –",
                    "Physics": "Yearly exam: समय - 							वार्षिक परीक्षा\nPhysics text for Class 1\n\nThe note: नोट – सभी प्रश्न अनिवार्य है –"
                },
                "Class-2": ["Eng", "Hindi", "Science", "Social", "Math", "Drawing", "History", "Physics"],
                # Add more classes here...
            },
            "The note": "                                                                                   नोट – सभी प्रश्न अनिवार्य है –"
        }
        self.create_widgets()
        self.preview_window = None
        self.recordings = []

    def create_widgets(self):
        # Menu
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.quit_app)
        menu.add_cascade(label="File", menu=file_menu)

        text_menu = tk.Menu(menu, tearoff=0)
        text_menu.add_command(label="Text 1", command=lambda: self.insert_predefined_text("text1.txt"))
        text_menu.add_command(label="Text 2", command=lambda: self.insert_predefined_text("text2.txt"))
        text_menu.add_command(label="Text 3", command=lambda: self.insert_predefined_text("text3.txt"))
        menu.add_cascade(label="MyText", menu=text_menu)

        class_menu = tk.Menu(menu, tearoff=0)
        for class_name, subjects in self.predefined_texts["Class"].items():
            sub_menu = tk.Menu(class_menu, tearoff=0)
            if isinstance(subjects, dict):
                for subject in subjects:
                    sub_menu.add_command(label=subject, command=lambda cls=class_name, sub=subject: self.insert_predefined_text2(cls, sub))
            else:
                sub_menu.add_command(label="Select", state=tk.DISABLED)
            class_menu.add_cascade(label=class_name, menu=sub_menu)
        menu.add_cascade(label="Class", menu=class_menu)

        text_menu2 = tk.Menu(menu, tearoff=0)
        for label in self.predefined_texts:
            text_menu2.add_command(label=label, command=lambda txt=self.predefined_texts[label]: self.insert_predefined_text2(txt))
        menu.add_cascade(label="Predefined Text", menu=text_menu2)        

        # Toolbar
        toolbar = tk.Frame(self.root, bg="#f0f0f0")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_button = tk.Button(toolbar, text="Start Recording", command=self.start_recording, state="normal")
        self.recording_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_period_button = tk.Button(toolbar, text="Punctuation: Period", command=lambda: self.set_punctuation("."))
        punctuation_period_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_question_button = tk.Button(toolbar, text="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))
        punctuation_question_button.pack(side=tk.LEFT, padx=5, pady=5)

        predefined_button = tk.Button(toolbar, text="Predefined Text", command=self.insert_predefined_text2)
        predefined_button.pack(side=tk.LEFT, padx=5, pady=5)

        remove_last_button = tk.Button(toolbar, text="Remove number", command=self.remove_number)
        remove_last_button.pack(side=tk.LEFT, padx=5, pady=5)

        copy_button = tk.Button(toolbar, text="Copy All", command=self.copy_text)
        copy_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.LEFT, padx=5, pady=5)

        # Text box
        self.text_box = tk.Text(self.root, font=("Arial", 14), wrap="word")
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Bind right-click menu
        self.text_box.bind("<Button-3>", self.show_context_menu)

    def insert_predefined_text2(self, class_name=None, subject=None):
        if class_name and subject:
            text = self.predefined_texts["Class"][class_name][subject]
        elif class_name:
            text = f"कक्षा - {class_name}"
        elif subject:
            text = f"विषय - {subject}"
        else:
            text = self.predefined_texts["The note"]
        self.text_box.insert(tk.END, text + "\n")


    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Select All", command=self.select_all)
        context_menu.add_separator()
        context_menu.add_command(label="Copy", command=self.copy_selected_text)
        context_menu.add_command(label="Cut", command=self.cut_selected_text)
        context_menu.add_command(label="Paste", command=self.paste_text)
        context_menu.tk_popup(event.x_root, event.y_root)

    def copy_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(text_content)
        messagebox.showinfo("Success", "Text copied successfully!")

    def save_to_docx(self):
        if self.recordings:
            print("Type of self.recordings:", type(self.recordings))  # Debugging print statement
            print("Content of self.recordings:", self.recordings)  # Debugging print statement
            current_dir = os.path.dirname(os.path.abspath(__file__))
            file_path = os.path.join(current_dir, "speech_to_text.docx")
            try:
                doc = Document(file_path) if os.path.exists(file_path) else Document()
                for text in self.recordings:
                    doc.add_paragraph(text)
                doc.save(file_path)
                self.recordings.clear()
            except Exception as e:
                print("Error saving file:", e)

    def start_recording(self):
        if not self.recording:
            self.recording = True
            self.recording_button.config(text="Stop Recording")
            self.status_label.config(text="Recording...", fg="red")
            threading.Thread(target=self._recording_thread).start()
        else:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread(self):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.hindi_numbers[self.counter - 1]}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            # Insert voice-recognized text without bold formatting
            self.text_box.insert(tk.END, numbered_text + "\n")
            self.recordings.append(numbered_text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")


    def set_punctuation(self, punctuation):
        self.selected_punctuation = punctuation

    def insert_predefined_text(self, filename=None):
        self.counter = 1  # Reset the counter to 1
        try:
            if filename:
                file_path = os.path.join(os.path.dirname(__file__), filename)
                print("File path:", file_path)  # Print the absolute file path for debugging
                with open(file_path, "r", encoding="utf-8") as file:
                    predefined_text = file.read()
                    numbered_text = ""
                    line = predefined_text.strip()
                    while line:
                        numbered_text += f"{self.roman_numbers[self.counter2 - 1]}. {line}\n\n"
                        self.counter2 += 1
                        line = file.readline().strip()
                    self.text_box.insert(tk.END, numbered_text)  # Insert predefined text
                    self.recordings = [numbered_text]  # Reset the recordings list
        except FileNotFoundError:
            print("Predefined text file not found.")

    def save_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        default_file_name = "My Exam Paper"  # Default file name without extension
        file_name = simpledialog.askstring("Save As", "Enter file name (without extension):", initialvalue=default_file_name)
        if file_name:
            file_name += ".docx"  # Append .docx extension
            file_path = os.path.join(os.getcwd(), file_name)
            try:
                doc = Document()  # Create a new Document object

                # Set page size to A4
                section = doc.sections[0]
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)

                # Set margins to narrow
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Define paragraph style with minimum spacing
                style = doc.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                # Add content to the document
                for line in text_content.split("\n"):
                    if line.strip():
                        paragraph = doc.add_paragraph(line.strip())

                doc.save(file_path)
                messagebox.showinfo("Success", f"File saved as {file_name}")
            except Exception as e:
                messagebox.showerror("Error", f"Error saving file: {e}")

    def print_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        try:
            # Create a new Document object
            doc = Document()

            # Add content to the document
            for line in text_content.split("\n"):
                if line.strip():
                    paragraph = doc.add_paragraph(line.strip())

            # Create a handle to the default printer
            printer_handle = win32print.GetDefaultPrinter()

            # Create a Device Context (DC) for the printer
            printer_dc = win32ui.CreateDC()

            # Set the DC to the printer
            printer_dc.CreatePrinterDC(printer_handle)

            # Start a print job
            printer_dc.StartDoc("My Document")

            # Start a new page
            printer_dc.StartPage()

            # Render the document to the printer
            doc._element = doc._body
            doc._saved = True
            doc._preamble = None
            doc.save(printer_dc)

            # End the print job
            printer_dc.EndPage()
            printer_dc.EndDoc()

            # Cleanup
            printer_dc.DeleteDC()

            messagebox.showinfo("Success", "Document sent to printer successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Error printing document: {e}")

    def preview_text(self):
        if self.preview_window:
            self.preview_window.destroy()

        self.preview_window = tk.Toplevel(self.root)
        self.preview_window.title("Preview")
        self.preview_window.geometry("600x400")

        preview_text_box = tk.Text(self.preview_window, font=("Arial", 14), wrap="word")
        preview_text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        preview_text_box.insert(tk.END, self.text_box.get("1.0", tk.END))
        preview_text_box.config(state=tk.DISABLED)

    def remove_number(self):
        self.counter -= 1
        self.counter2 -= 1

    def select_all(self):
        self.text_box.tag_add(tk.SEL, "1.0", tk.END)
        self.text_box.mark_set(tk.INSERT, "1.0")
        self.text_box.see(tk.INSERT)
        return 'break'

    def copy_selected_text(self):
        self.text_box.clipboard_clear()
        text = self.text_box.get(tk.SEL_FIRST, tk.SEL_LAST)
        self.text_box.clipboard_append(text)

    def cut_selected_text(self):
        self.copy_selected_text()
        self.text_box.delete(tk.SEL_FIRST, tk.SEL_LAST)

    def paste_text(self):
        text = self.text_box.clipboard_get()
        self.text_box.insert(tk.INSERT, text)

    def quit_app(self):
        self.save_to_docx()
        self.root.quit()


root = tk.Tk()
app = SpeechToTextApp(root)
root.mainloop()
