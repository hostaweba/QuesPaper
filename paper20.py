import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import RGBColor
import win32print
import win32ui
import os
import threading
import time
import subprocess

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.root.geometry("800x600")
        self.root.resizable(False, False)
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.hindi_numbers = ['(क)', '(ख)', '(ग)', '(घ)', '(ङ)', '(च)', '(छ)', '(ज)', '(झ)', '(ञ)']  # Hindi numbers
        self.counter2 = 1  # Initialize counter for numbering
        self.roman_numbers = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI', 'XVII', 'XVIII', 'XIX','XX']  # Hindi numbers
        self.selected_punctuation = "?"  # Default punctuation
        self.recording = False  # Flag to indicate recording status
        
        self.another_predefined_texts = {
            "Example 1": "This is the text for Example 1.",
            "Example 2": "This is the text for Example 2.",
            "Example 3": "This is the text for Example 3."
        }        
        
        self.predefined_texts = {
            "Yearly exam": "समय - 						    वार्षिक परीक्षा",
            "Half Yearly exam": "समय - 						    अर्ध-वार्षिक परीक्षा",
            "The note": "                                                                                   नोट – सभी प्रश्न अनिवार्य है –"
        }
        
        self.my_predefined_texts = {
            "Class-1": {
                "Hindi": "                                                                                     कक्षा – 1 विषय – हिन्दी",
                "Social": "                                                                                     कक्षा – 1 विषय – सामाजिक",
                "Science": "                                                                                     कक्षा – 1 विषय – विज्ञान"
            },
            "Class-2": {
                "The Architect": "Savita - The Architect",
                "The Doctor": "Savita - The Doctor",
                "The Teacher": "Savita - The Teacher"
            },
            "Class-3": {
                "The Architect": "Mohini - The Architect",
                "The Doctor": "Mohini - The Doctor",
                "The Teacher": "Mohini - The Teacher"
            }
        } 

 
        self.create_widgets()
        self.preview_window = None
        self.recordings = []

    def create_widgets(self):
        # Menu
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Open Saved File", command=self.open_saved_file)  # Add command to open saved file
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.quit_app)
        menu.add_cascade(label="File", menu=file_menu)




        text_menu = tk.Menu(menu, tearoff=0)
        text_menu.add_command(label="Text 1", command=lambda: self.insert_predefined_text("text1.txt"))
        text_menu.add_command(label="Text 2", command=lambda: self.insert_predefined_text("text2.txt"))
        text_menu.add_command(label="Text 3", command=lambda: self.insert_predefined_text("text3.txt"))
        menu.add_cascade(label="MyText", menu=text_menu)
        
        another_predefined_submenu = tk.Menu(menu, tearoff=0)
        for label, text in self.another_predefined_texts.items():
            another_predefined_submenu.add_command(label=label, command=lambda txt=text: self.insert_predefined_text2(txt))
        menu.add_cascade(label="Another Predefined Text", menu=another_predefined_submenu)       
        

        text_menu2 = tk.Menu(menu, tearoff=0)
        for label in self.predefined_texts:
            text_menu2.add_command(label=label, command=lambda txt=self.predefined_texts[label]: self.insert_predefined_text2(txt))
        menu.add_cascade(label="Exam Title", menu=text_menu2)        



        # "many text" submenu
        many_text_submenu = tk.Menu(menu, tearoff=0)
        for name, texts in self.my_predefined_texts.items():
            submenu = tk.Menu(many_text_submenu, tearoff=0)
            for label, text in texts.items():
                submenu.add_command(label=label, command=lambda txt=text: self.insert_predefined_text2(txt))
            many_text_submenu.add_cascade(label=name, menu=submenu)
        menu.add_cascade(label="Class & SubjectS", menu=many_text_submenu)

        # Toolbar
        toolbar = tk.Frame(self.root, bg="#f0f0f0")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_button = tk.Button(toolbar, text="Start Recording", command=self.start_recording, state="normal")
        self.recording_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_period_button = tk.Button(toolbar, text="Punctuation: Period", command=lambda: self.set_punctuation("."))
        punctuation_period_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_question_button = tk.Button(toolbar, text="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))
        punctuation_question_button.pack(side=tk.LEFT, padx=5, pady=5)

        predefined_button = tk.Button(toolbar, text="Predefined Text", command=self.insert_predefined_text2)
        predefined_button.pack(side=tk.LEFT, padx=5, pady=5)

        remove_last_button = tk.Button(toolbar, text="decrease number", command=self.remove_number)
        remove_last_button.pack(side=tk.LEFT, padx=5, pady=5)

        copy_button = tk.Button(toolbar, text="Copy All", command=self.copy_text)
        copy_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.LEFT, padx=5, pady=5)

        # Text box
        self.text_box = tk.Text(self.root, font=("Arial", 14), wrap="word")
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Bind right-click menu
        self.text_box.bind("<Button-3>", self.show_context_menu)

    def insert_predefined_text2(self, text=None):
        if text is None:
            text = self.predefined_texts["Text 1"]  # Default to the first predefined text
        
        # Apply bold formatting to the predefined text
        self.text_box.insert(tk.END, text + "\n", "predefined")  # Append newline character
        self.text_box.tag_configure("predefined", font=("Arial", 14, "bold"))



    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Select All", command=self.select_all)
        context_menu.add_separator()
        context_menu.add_command(label="Copy", command=self.copy_selected_text)
        context_menu.add_command(label="Cut", command=self.cut_selected_text)
        context_menu.add_command(label="Paste", command=self.paste_text)
        context_menu.tk_popup(event.x_root, event.y_root)

    def copy_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(text_content)
        messagebox.showinfo("Success", "Text copied successfully!")

    def save_to_docx(self):
        if self.recordings:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            file_path = os.path.join(current_dir, "speech_to_text.docx")
            try:
                doc = Document(file_path) if os.path.exists(file_path) else Document()
                for text in self.recordings:
                    doc.add_paragraph(text)
                doc.save(file_path)
                self.recordings.clear()
            except Exception as e:
                print("Error saving file:", e)

    def start_recording(self):
        if not self.recording:
            self.recording = True
            self.recording_button.config(text="Stop Recording")
            self.status_label.config(text="Recording...", fg="red")
            threading.Thread(target=self._recording_thread).start()
        else:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread(self):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.hindi_numbers[self.counter - 1]}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            # Insert voice-recognized text without bold formatting
            self.text_box.insert(tk.END, numbered_text + "\n")
            self.recordings.append(numbered_text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")


    def set_punctuation(self, punctuation):
        self.selected_punctuation = punctuation

    def insert_predefined_text(self, filename=None):
        self.counter = 1  # Reset the counter to 1
        try:
            if filename:
                file_path = os.path.join(os.path.dirname(__file__), filename)
                print("File path:", file_path)  # Print the absolute file path for debugging
                with open(file_path, "r", encoding="utf-8") as file:
                    predefined_text = file.read()
                    numbered_text = ""
                    line = predefined_text.strip()
                    while line:
                        numbered_text += f"{self.roman_numbers[self.counter2 - 1]}. {line}\n\n"
                        self.counter2 += 1
                        line = file.readline().strip()
                    self.text_box.insert(tk.END, numbered_text)  # Insert predefined text
                    self.recordings = [numbered_text]  # Reset the recordings list
        except FileNotFoundError:
            print("Predefined text file not found.")

 
    def save_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        default_file_name = "My Exam Paper"  # Default file name without extension
        file_name = simpledialog.askstring("Save As", "Enter file name (without extension):", initialvalue=default_file_name)
        if file_name:
            file_name += ".docx"  # Append .docx extension
            file_path = os.path.join(os.getcwd(), file_name)
            try:
                doc = Document()  # Create a new Document object

                # Set page size to A4
                section = doc.sections[0]
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)

                # Set margins to narrow
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Define paragraph style with minimum spacing
                style = doc.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                # Iterate over each line in the text content
                for line in text_content.split("\n"):
                    # Check if the line contains a predefined text
                    predefined_text_found = False
                    
                    # Check in the first dictionary
                    for predefined_text in self.predefined_texts.values():
                        if predefined_text in line:
                            # Apply bold formatting to the line
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run(line)
                            run.font.bold = True
                            run.font.size = Pt(12)  # Set font size to 14 points
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                            predefined_text_found = True
                            break
                    
                    # If not found in the first dictionary, check in the second dictionary
                    if not predefined_text_found:
                        for predefined_text in self.another_predefined_texts.values():
                            if predefined_text in line:
                                # Apply bold formatting to the line
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(line)
                                run.font.bold = True
                                run.font.size = Pt(12)  # Set font size to 14 points
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                                predefined_text_found = True
                                break
                    
                    # If not found in the second dictionary, check in the third dictionary
                    if not predefined_text_found:
                        for person, person_texts in self.my_predefined_texts.items():
                            for text_name, text_content in person_texts.items():
                                if text_content in line:
                                    # Apply bold formatting to the line
                                    paragraph = doc.add_paragraph()
                                    run = paragraph.add_run(line)
                                    run.font.bold = True
                                    run.font.size = Pt(12)  # Set font size to 14 points
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                                    predefined_text_found = True
                                    break
                            if predefined_text_found:
                                break
                    
                    if not predefined_text_found:
                        # If the line does not contain a predefined text, add it normally
                        paragraph = doc.add_paragraph(line, style)
                        # Apply font type and size to the entire paragraph
                        for run in paragraph.runs:
                            run.font.name = "Arial"  # Set font type to Arial
                            run.font.size = Pt(11)  # Set font size to 14 points

                doc.save(file_path)  # Save the document
                messagebox.showinfo("Success", f"Text saved successfully as '{file_name}'!")
            except PermissionError as e:
                messagebox.showerror("Error", f"Please close the {file_name} file before saving.")
            except Exception as e:
                print("Error saving file:", e)




    def open_saved_file(self):
        current_dir = os.getcwd()
        file_path = filedialog.askopenfilename(initialdir=current_dir, title="Select file", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
        if file_path:
            try:
                os.startfile(file_path)  # Open the file using default application
            except Exception as e:
                print("Error opening file:", e)



    def preview_text(self):
        if self.preview_window is None or not self.preview_window.winfo_exists():
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title("Preview")
            preview_textbox = tk.Text(self.preview_window, font=("Arial", 14), wrap="word")
            preview_textbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            preview_text = self.text_box.get("1.0", tk.END)
            preview_textbox.insert(tk.END, preview_text)
        else:
            self.preview_window.deiconify()
            self.preview_window.lift()

    def print_text(self):
        text = self.text_box.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)

    def remove_number(self):
        if self.counter > 1:
            # Find the start position of the last line number
            start_pos = self.text_box.search(f"\\(.*?\\)", "1.0", stopindex=tk.END, regexp=True)
            print("Start position:", start_pos)  # Debugging print statement
            if start_pos:
                # Find the end position of the last line
                end_pos = f"{self.hindi_numbers[self.counter - 2]}.999"  # Use 999 as the column number to get the end of the line
                print("End position:", end_pos)  # Debugging print statement

                self.counter -= 1
                self.text_box.delete(start_pos, end_pos)

    def quit_app(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()

if __name__ == "__main__":
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
