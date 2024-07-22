import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
import speech_recognition as sr
import win32print
import win32ui
import os
import threading

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.root.geometry("800x600")
        self.root.resizable(False, False)
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.hindi_numbers = ['(क)', '(ख)', '(ग)', '(घ)', '(ङ)', '(च)', '(छ)', '(ज)', '(झ)', '(ञ)', '(ट)', '(ठ)', '(ड)', '(ढ)', '(ण)', '(त)', '(थ)', '(द)', '(ध)', '(न)', '(प)', '(फ)', '(ब)', '(भ)', '(म)', '(य)', '(र)', '(ल)', '(व)', '(श)', '(ष)', '(स)', '(ह)', '(अ)', '(आ)', '(इ)', '(ई)', '(उ)', '(ऊ)', '(ऋ)', '(ए)', '(ऐ)', '(ओ)', '(औ)', '(अं)', '(अः)', '(क्ष)', '(त्र)', '(ज्ञ)']  # Hindi numbers
        self.counter2 = 1  # Initialize counter for numbering
        self.roman_numbers = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI', 'XVII', 'XVIII', 'XIX','XX']  # Hindi numbers
        self.selected_punctuation = "?"  # Default punctuation
        self.recording = False  # Flag to indicate recording status
        self.recordings = []

        self.create_widgets()
        self.preview_window = None

    def create_widgets(self):
        # Menu
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Open Saved File", command=self.open_saved_file)
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Exit", command=self.quit_app)
        menu.add_cascade(label="File", menu=file_menu)

        text_menu = tk.Menu(menu, tearoff=0)
        text_menu.add_command(label="Add Multiple Choice", command=self.add_multiple_choice)
        menu.add_cascade(label="Text", menu=text_menu)

        # Toolbar
        toolbar = tk.Frame(self.root, bg="#f0f0f0")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_button = tk.Button(toolbar, text="Start Recording", command=self.start_recording, state="normal")
        self.recording_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_period_button = tk.Button(toolbar, text="Punctuation: Period", command=lambda: self.set_punctuation("|"))
        punctuation_period_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_question_button = tk.Button(toolbar, text="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))
        punctuation_question_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.LEFT, padx=5, pady=5)

        # Text box
        self.text_box = tk.Text(self.root, font=("Arial", 14), wrap="word")
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Bind right-click menu
        self.text_box.bind("<Button-3>", self.show_context_menu)

    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Select All", command=self.select_all)
        context_menu.add_separator()
        context_menu.add_command(label="Copy", command=self.copy_selected_text)
        context_menu.add_command(label="Cut", command=self.cut_selected_text)
        context_menu.add_command(label="Paste", command=self.paste_text)
        context_menu.tk_popup(event.x_root, event.y_root)

    def start_recording(self):
        if not self.recording:
            self.recording = True
            self.recording_button.config(text="Stop Recording")
            self.status_label.config(text="Recording...", fg="red")
            threading.Thread(target=self._recording_thread).start()
        else:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def _recording_thread(self):
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.hindi_numbers[self.counter - 1]}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            # Insert voice-recognized text without bold formatting
            self.text_box.insert(tk.END, numbered_text + "\n")
            self.recordings.append(numbered_text)  # Save for batch saving
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
        except Exception as e:
            print("Error:", e)
        finally:
            self.recording = False
            self.recording_button.config(text="Start Recording")
            self.status_label.config(text="Ready", fg="green")

    def set_punctuation(self, punctuation):
        self.selected_punctuation = punctuation

    def save_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        default_file_name = "My Exam Paper"  # Default file name without extension
        file_name = simpledialog.askstring("Save As", "Enter file name (without extension):", initialvalue=default_file_name)
        if file_name:
            file_name += ".docx"  # Append .docx extension
            file_path = os.path.join(os.getcwd(), file_name)
            try:
                doc = Document()  # Create a new Document object

                # Set page size to A4
                section = doc.sections[0]
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)

                # Set margins to narrow
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Define paragraph style with minimum spacing
                style = doc.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                # Iterate over each line in the text content
                for line in text_content.split("\n"):
                    # Check if the line contains a predefined text
                    predefined_text_found = False

                    if predefined_text_found:
                        # Apply bold formatting to the entire paragraph
                        paragraph = doc.add_paragraph(line)
                        paragraph.style = 'Strong'
                    else:
                        # If the line does not contain a predefined text, add it normally
                        paragraph = doc.add_paragraph(line, style)
                        # Apply font type and size to the entire paragraph
                        for run in paragraph.runs:
                            run.font.name = "Arial"  # Set font type to Arial
                            run.font.size = Pt(11)  # Set font size to 14 points

                doc.save(file_path)  # Save the document
                messagebox.showinfo("Success", f"Text saved successfully as '{file_name}'!")
            except PermissionError as e:
                messagebox.showerror("Error", f"Please close the {file_name} file before saving.")
            except Exception as e:
                print("Error saving file:", e)

    def open_saved_file(self):
        current_dir = os.getcwd()
        file_path = filedialog.askopenfilename(initialdir=current_dir, title="Select file", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
        if file_path:
            try:
                os.startfile(file_path)  # Open the file using default application
            except Exception as e:
                print("Error opening file:", e)

    def print_text(self):
        text = self.text_box.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)

    def add_multiple_choice(self):
        # Start speech recognition for question
        with self.microphone as source:
            print("Speak your question:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            question_audio = self.recognizer.listen(source)
        
        try:
            question = self.recognizer.recognize_google(question_audio)
            print("Question:", question)
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand the question")
            return
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
            return

        # Prompt user for choices
        print("Speak your choices separated by 'and':")
        with self.microphone as source:
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            choices_audio = self.recognizer.listen(source)
        
        try:
            choices = self.recognizer.recognize_google(choices_audio)
            print("Choices:", choices)
            choices = choices.split(" and ")
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand the choices")
            return
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
            return

        # Append the question to the text box
        self.text_box.insert(tk.END, f"\n{self.roman_numbers[self.counter2 - 1]}. {question}\n")
        
        # Append the choices to the text box
        for index, choice in enumerate(choices, start=1):
            self.text_box.insert(tk.END, f"{chr(96 + index)}. {choice.strip()}\n")
        
        # Increment the counter for numbering
        self.counter2 += 1

        # Apply bold formatting to the entire inserted text
        self.text_box.tag_configure("bold", font=("Arial", 14, "bold"))
        self.text_box.tag_add("bold", "end - 2 lines", "end - 1 lines")

        # Update recordings list
        self.recordings.append(question)

        # Add choices to the recordings list
        self.recordings.extend(choices)

    def quit_app(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()

if __name__ == "__main__":
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
