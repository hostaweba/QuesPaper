import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Pt, Inches, Cm
import win32print
import win32ui
import os
import threading

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.root.geometry("800x600")
        self.root.resizable(False, False)
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.counter = 1  # Initialize counter for numbering
        self.hindi_numbers = ['(क)', '(ख)', '(ग)', '(घ)', '(ङ)', '(च)', '(छ)', '(ज)', '(झ)', '(ञ)']  # Hindi numbers
        self.counter2 = 1  # Initialize counter for numbering
        self.roman_numbers = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI', 'XVII', 'XVIII', 'XIX','XX']  # Hindi numbers
        self.selected_punctuation = "?"  # Default punctuation
        self.recording = False  # Flag to indicate recording status
        self.predefined_texts = {
            "Yearly exam": "समय - 						    वार्षिक परीक्षा",
            "Class": {
                "Class-1": ["Eng", "Hindi", "Science", "Social", "Math", "Drawing", "History", "Physics"],
                "Class-2": ["Eng", "Hindi", "Science", "Social", "Math", "Drawing", "History", "Physics"],
                # Add more classes here...
            },
            "The note": "                                                                                   नोट – सभी प्रश्न अनिवार्य है –"
        }
        self.create_widgets()
        self.preview_window = None
        self.recordings = []

    def create_widgets(self):
        # Menu
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.quit_app)
        menu.add_cascade(label="File", menu=file_menu)

        text_menu = tk.Menu(menu, tearoff=0)
        text_menu.add_command(label="Text 1", command=lambda: self.insert_predefined_text("text1.txt"))
        text_menu.add_command(label="Text 2", command=lambda: self.insert_predefined_text("text2.txt"))
        text_menu.add_command(label="Text 3", command=lambda: self.insert_predefined_text("text3.txt"))
        menu.add_cascade(label="MyText", menu=text_menu)

        class_menu = tk.Menu(menu, tearoff=0)
        for class_name, subjects in self.predefined_texts["Class"].items():
            sub_menu = tk.Menu(class_menu, tearoff=0)
            for subject in subjects:
                sub_menu.add_command(label=subject, command=lambda cls=class_name, sub=subject: self.insert_predefined_text2(cls + ", " + sub))
            class_menu.add_cascade(label=class_name, menu=sub_menu)
        menu.add_cascade(label="Class", menu=class_menu)

        text_menu2 = tk.Menu(menu, tearoff=0)
        for label in self.predefined_texts:
            text_menu2.add_command(label=label, command=lambda txt=self.predefined_texts[label]: self.insert_predefined_text2(txt))
        menu.add_cascade(label="Predefined Text", menu=text_menu2)        

        # Toolbar
        toolbar = tk.Frame(self.root, bg="#f0f0f0")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.recording_button = tk.Button(toolbar, text="Start Recording", command=self.start_recording, state="normal")
        self.recording_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_period_button = tk.Button(toolbar, text="Punctuation: Period", command=lambda: self.set_punctuation("."))
        punctuation_period_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_question_button = tk.Button(toolbar, text="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))
        punctuation_question_button.pack(side=tk.LEFT, padx=5, pady=5)

        predefined_button = tk.Button(toolbar, text="Predefined Text", command=self.insert_predefined_text2)
        predefined_button.pack(side=tk.LEFT, padx=5, pady=5)

        remove_last_button = tk.Button(toolbar, text="decrease number", command=self.remove_number)
        remove_last_button.pack(side=tk.LEFT, padx=5, pady=5)

        copy_button = tk.Button(toolbar, text="Copy All", command=self.copy_text)
        copy_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.status_label = tk.Label(toolbar, text="Ready", fg="green")
        self.status_label.pack(side=tk.LEFT, padx=5, pady=5)

        # Text box
        self.text_box = tk.Text(self.root, font=("Arial", 14), wrap="word")
        self.text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Bind right-click menu
        self.text_box.bind("<Button-3>", self.show_context_menu)

    def start_recording(self):
        pass  # Placeholder for the start_recording method

    def insert_predefined_text2(self, text=None):
        if text is None:
            text = self.predefined_texts["The note"]
        # Apply bold formatting to the predefined text
        self.text_box.insert(tk.END, text + "\n", "predefined")  # Append newline character
        self.text_box.tag_configure("predefined", font=("Arial", 14, "bold"))

    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Select All", command=self.select_all)
        context_menu.add_separator()
        context_menu.add_command(label="Copy", command=self.copy_selected_text)
        context_menu.add_command(label="Cut", command=self.cut_selected_text)
        context_menu.add_command(label="Paste", command=self.paste_text)
        context_menu.tk_popup(event.x_root, event.y_root)

    def copy_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(text_content)
        messagebox.showinfo("Success", "Text copied successfully!")

    def save_text(self):
        text_content = self.text_box.get("1.0", tk.END)
        default_file_name = "My Exam Paper"  # Default file name without extension
        file_name = simpledialog.askstring("Save As", "Enter file name (without extension):", initialvalue=default_file_name)
        if file_name:
            file_name += ".docx"  # Append .docx extension
            file_path = os.path.join(os.getcwd(), file_name)
            try:
                doc = Document()  # Create a new Document object

                # Set page size to A4
                section = doc.sections[0]
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)

                # Set margins to narrow
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                # Define paragraph style with minimum spacing
                style = doc.styles['Normal']
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

                # Iterate over each line in the text content
                for line in text_content.split("\n"):
                    # Check if the line contains a predefined text
                    predefined_text_found = False
                    for predefined_text in self.predefined_texts.values():
                        if predefined_text in line:
                            # Apply bold formatting to the line
                            paragraph = doc.add_paragraph(line)
                            paragraph.style = style
                            run = paragraph.runs[0]
                            run.bold = True
                            predefined_text_found = True
                            break
                    if not predefined_text_found:
                        # If the line does not contain a predefined text, add it normally
                        paragraph = doc.add_paragraph(line)
                        paragraph.style = style

                doc.save(file_path)  # Save the document
                messagebox.showinfo("Success", f"Text saved successfully as '{file_name}'!")
            except PermissionError as e:
                messagebox.showerror("Error", f"Please close the {file_name} file before saving.")
            except Exception as e:
                print("Error saving file:", e)

    def preview_text(self):
        if self.preview_window is None or not self.preview_window.winfo_exists():
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title("Preview")
            preview_textbox = tk.Text(self.preview_window, font=("Arial", 14), wrap="word")
            preview_textbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            preview_text = self.text_box.get("1.0", tk.END)
            preview_textbox.insert(tk.END, preview_text)
        else:
            self.preview_window.deiconify()
            self.preview_window.lift()

    def print_text(self):
        text = self.text_box.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)

    def remove_number(self):
        if self.counter > 1:
            # Find the start position of the last line number
            start_pos = self.text_box.search(f"\\(.*?\\)", "1.0", stopindex=tk.END, regexp=True)
            print("Start position:", start_pos)  # Debugging print statement
            if start_pos:
                # Find the end position of the last line
                end_pos = f"{self.hindi_numbers[self.counter - 2]}.999"  # Use 999 as the column number to get the end of the line
                print("End position:", end_pos)  # Debugging print statement

                self.counter -= 1
                self.text_box.delete(start_pos, end_pos)

    def quit_app(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()

if __name__ == "__main__":
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
