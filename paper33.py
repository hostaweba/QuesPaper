import speech_recognition as sr
import tkinter as tk
from tkinter import filedialog
from docx import Document
import win32print
import win32ui
import os

class SpeechToTextApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text App")
        self.recognizer = None
        self.microphone = None
        self.counter = 1  # Initialize counter for numbering
        self.selected_punctuation = "."  # Default punctuation
        self.conversion_enabled = False  # Flag to indicate if English to Hindi conversion is enabled
        self.create_menu()
        self.create_toolbar()
        self.create_textbox()
        self.preview_window = None

    def create_menu(self):
        menubar = tk.Menu(self.root)
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Save", command=self.save_text)
        file_menu.add_command(label="Print", command=self.print_text)
        file_menu.add_command(label="Preview", command=self.preview_text)
        file_menu.add_command(label="Exit", command=self.quit_app)
        menubar.add_cascade(label="File", menu=file_menu)
        self.root.config(menu=menubar)

    def create_toolbar(self):
        toolbar = tk.Frame(self.root, bg="lightgray")
        toolbar.pack(side=tk.TOP, fill=tk.X)

        start_button = tk.Button(toolbar, text="Start Recording", command=self.start_recording)
        start_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_period_button = tk.Button(toolbar, text="Punctuation: Period", command=lambda: self.set_punctuation("."))
        punctuation_period_button.pack(side=tk.LEFT, padx=5, pady=5)

        punctuation_question_button = tk.Button(toolbar, text="Punctuation: Question Mark", command=lambda: self.set_punctuation("?"))
        punctuation_question_button.pack(side=tk.LEFT, padx=5, pady=5)

        predefined_button = tk.Button(toolbar, text="Predefined Text", command=self.insert_predefined_text)
        predefined_button.pack(side=tk.LEFT, padx=5, pady=5)

        toggle_conversion_button = tk.Button(toolbar, text="Toggle Conversion", command=self.toggle_conversion)
        toggle_conversion_button.pack(side=tk.LEFT, padx=5, pady=5)

    def create_textbox(self):
        self.textbox = tk.Text(self.root, height=20, width=60)
        self.textbox.pack(pady=10)

    def save_to_docx(self, text):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(current_dir, "speech_to_text.docx")
        try:
            doc = Document(file_path)
        except FileNotFoundError:
            doc = Document()
        doc.add_paragraph(text)
        doc.save(file_path)

    def start_recording(self):
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        with self.microphone as source:
            print("Speak in Hindi:")
            self.recognizer.adjust_for_ambient_noise(source, duration=1)
            self.audio = self.recognizer.listen(source)

        try:
            text = self.recognizer.recognize_google(self.audio, language="hi-IN")
            print("You said:", text)

            # Append counter to the recognized text
            numbered_text = f"{self.counter}. {text}{self.selected_punctuation}"  # Append selected punctuation
            self.counter += 1  # Increment counter

            if self.conversion_enabled:
                numbered_text = self.convert_to_hindi(numbered_text)

            self.textbox.insert(tk.END, numbered_text + "\n")  # Display in text box
            self.save_to_docx(numbered_text)  # Save to document
            
        except sr.UnknownValueError:
            print("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
        except Exception as e:
            print("Error:", e)

    def convert_to_hindi(self, text):
        # Implement English to Hindi conversion logic here
        # This is just a placeholder function, replace it with actual conversion logic
        return text

    def set_punctuation(self, punctuation):
        self.selected_punctuation = punctuation

    def insert_predefined_text(self):
        self.counter = 1  # Reset the counter to 1
        try:
            with open("predefined_text.txt", "r", encoding="utf-8") as file:
                predefined_text = file.read()
                self.textbox.insert(tk.END, predefined_text + "\n")
        except FileNotFoundError:
            print("Predefined text file not found.")

    def toggle_conversion(self):
        self.conversion_enabled = not self.conversion_enabled

    def save_text(self):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = filedialog.asksaveasfilename(initialdir=current_dir, defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if file_path:
            text = self.textbox.get("1.0", tk.END)
            try:
                doc = Document()
                doc.add_paragraph(text)
                doc.save(file_path)
            except Exception as e:
                print("Error saving file:", e)

    def preview_text(self):
        if self.preview_window is None or not self.preview_window.winfo_exists():
            self.preview_window = tk.Toplevel(self.root)
            self.preview_window.title("Preview - Print Layout")

            # Create a canvas to simulate an A4 page layout (595x842 points)
            canvas = tk.Canvas(self.preview_window, width=595, height=842, bg="white")
            canvas.pack()

            # Add margins and text to the canvas
            margin = 50
            text = self.textbox.get("1.0", tk.END)
            canvas.create_text(margin, margin, anchor=tk.NW, text=text, font=("Courier", 12), width=595 - 2 * margin)
        else:
            self.preview_window.deiconify()
            self.preview_window.lift()

    def print_text(self):
        text = self.textbox.get("1.0", tk.END)
        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        hdc = win32ui.CreateDC()
        hdc.CreatePrinterDC(printer_name)
        hdc.StartDoc('Speech to Text Document')
        hdc.StartPage()
        hdc.DrawText(text, (100, 100, 1000, 1000), win32ui.DT_LEFT)
        hdc.EndPage()
        hdc.EndDoc()
        hdc.DeleteDC()
        win32print.ClosePrinter(hprinter)

    def quit_app(self):
        self.root.quit()

if __name__ == "__main__":
    root = tk.Tk()
    app = SpeechToTextApp(root)
    root.mainloop()
